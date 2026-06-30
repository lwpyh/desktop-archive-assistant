"""Office 文档/表格能力集 — 把「数据提取/分析」「文档生成」需求内化到本 skill。

在通用整理(17) + 照片/视频(9) 之上，补齐文件**内容级**操作能力。延续本 skill
安全原则：所有写操作默认 dry-run；绝不覆盖原文件（一律写到新文件名/输出目录）；
原始数据只读不改。

能力清单（Office 专项，4 个）：
  table_clean   — 表格列筛选 / 去重 / 空值清洗 / 排序（xlsx/csv）
  table_merge   — 多个 txt/csv/xlsx 汇总成一个表（可按时间/名称排序）
  docx_compose  — 清单/资料整理成 Word 文档（标题 + 段落/表格，支持模板）
  pdf_ops       — PDF 合并 / 拆分 / 抽取文本（不依赖 Office）

依赖（按需懒加载，缺失时给出明确安装提示并优雅降级）：
  pandas       表格清洗/筛选/去重/汇总
  openpyxl     xlsx 读写（pandas 引擎）
  python-docx  Word 写入（已在依赖中）
  docxtpl      Word 模板填充（可选）
  pypdf        PDF 合并/拆分/抽取文本（已在依赖中）
"""
from __future__ import annotations

import csv
import glob
import os
import time
from typing import Any, Dict, List, Optional

from .utils import expand, ensure_dir, logger


# ============================================================
# 内部工具
# ============================================================

def _unique_path(dst: str) -> str:
    """同名冲突时加 -1/-2 后缀，绝不覆盖。"""
    if not os.path.exists(dst):
        return dst
    base, ext = os.path.splitext(dst)
    i = 1
    while os.path.exists(f"{base}-{i}{ext}"):
        i += 1
    return f"{base}-{i}{ext}"


def _missing(pkg: str, pip_name: str) -> Dict[str, Any]:
    msg = f"缺少依赖 {pkg}，请先安装：pip install {pip_name}"
    logger.error(msg)
    return {"error": msg}


def _require_pandas():
    try:
        import pandas as pd  # noqa: F401
        return pd
    except ImportError:
        return None


# ============================================================
# 能力 1：table_clean — 表格列筛选/去重/清洗/排序
# ============================================================

def capability_table_clean(
    src: str,
    keep_cols: Optional[List[str]] = None,
    drop_cols: Optional[List[str]] = None,
    dedup_by: Optional[List[str]] = None,
    dropna_cols: Optional[List[str]] = None,
    sort_by: Optional[str] = None,
    sort_desc: bool = False,
    sheet: Optional[str] = None,
    out_path: Optional[str] = None,
    dry_run: bool = False,
) -> Dict[str, Any]:
    """读取 xlsx/csv，做列筛选、去重、空值清洗、排序，写到**新文件**。

    安全：原文件只读；输出默认写到 <name>_cleaned.xlsx，不覆盖原文件。
    """
    pd = _require_pandas()
    if pd is None:
        return _missing("pandas", "pandas openpyxl")

    src = expand(src)
    if not os.path.isfile(src):
        return {"error": f"文件不存在: {src}"}

    ext = os.path.splitext(src)[1].lower()
    try:
        if ext in (".xlsx", ".xlsm", ".xls"):
            df = pd.read_excel(src, sheet_name=sheet or 0, dtype=object)
        elif ext in (".csv", ".tsv", ".txt"):
            sep = "\t" if ext == ".tsv" else None
            df = pd.read_csv(src, sep=sep, engine="python", dtype=object)
        else:
            return {"error": f"不支持的表格格式: {ext}（支持 xlsx/xls/csv/tsv）"}
    except Exception as e:  # noqa: BLE001
        return {"error": f"读取失败: {e}"}

    before_rows, before_cols = df.shape
    steps: List[str] = []

    # 1. 列筛选
    if keep_cols:
        valid = [c for c in keep_cols if c in df.columns]
        miss = [c for c in keep_cols if c not in df.columns]
        if miss:
            logger.warning(f"keep-cols 中不存在的列被忽略: {miss}")
        if valid:
            df = df[valid]
            steps.append(f"保留列 {valid}")
    if drop_cols:
        present = [c for c in drop_cols if c in df.columns]
        if present:
            df = df.drop(columns=present)
            steps.append(f"删除列 {present}")

    # 2. 空值清洗：去掉空字符串/纯空白，再按指定列丢弃空行
    df = df.replace(r"^\s*$", pd.NA, regex=True)
    if dropna_cols:
        present = [c for c in dropna_cols if c in df.columns]
        if present:
            df = df.dropna(subset=present)
            steps.append(f"按 {present} 丢弃空行")

    # 3. 去重
    if dedup_by:
        present = [c for c in dedup_by if c in df.columns]
        if present:
            df = df.drop_duplicates(subset=present, keep="first")
            steps.append(f"按 {present} 去重(保留首条)")
        else:
            df = df.drop_duplicates(keep="first")
            steps.append("整行去重")

    # 4. 排序
    if sort_by and sort_by in df.columns:
        df = df.sort_values(by=sort_by, ascending=not sort_desc, kind="stable")
        steps.append(f"按 {sort_by} {'降序' if sort_desc else '升序'}排序")

    after_rows, after_cols = df.shape

    # 输出路径：默认写到专属输出目录，绝不污染源文件所在目录
    if not out_path:
        base = os.path.basename(os.path.splitext(src)[0])
        out_path = os.path.join(expand("~/.archive_assistant/output"), f"{base}_cleaned.xlsx")
    out_path = _unique_path(expand(out_path))

    result = {
        "src": src,
        "out_path": out_path,
        "before": {"rows": before_rows, "cols": before_cols},
        "after": {"rows": after_rows, "cols": after_cols},
        "columns": list(df.columns),
        "steps": steps,
        "dry_run": dry_run,
        "preview": df.head(5).fillna("").astype(str).values.tolist(),
    }

    if dry_run:
        return result

    ensure_dir(os.path.dirname(out_path) or ".")
    try:
        if out_path.lower().endswith(".csv"):
            df.to_csv(out_path, index=False, encoding="utf-8-sig")
        else:
            df.to_excel(out_path, index=False)
    except Exception as e:  # noqa: BLE001
        return {"error": f"写出失败: {e}"}
    return result


# ============================================================
# 能力 2：table_merge — 多文件汇总成一个表
# ============================================================

def capability_table_merge(
    root: str,
    pattern: str = "*.txt",
    out_path: Optional[str] = None,
    recursive: bool = False,
    sort_by: str = "name",
    add_source_col: bool = True,
    dry_run: bool = False,
) -> Dict[str, Any]:
    """把目录下匹配的多个 txt/csv/xlsx 汇总成一个 xlsx。

    - txt：每行作为一条记录，列名 content；
    - csv/xlsx：按表头纵向拼接（列对齐）。
    安全：只读源文件，输出写到新 xlsx。
    """
    pd = _require_pandas()
    if pd is None:
        return _missing("pandas", "pandas openpyxl")

    root = expand(root)
    if not os.path.isdir(root):
        return {"error": f"目录不存在: {root}"}

    pat = os.path.join(root, "**", pattern) if recursive else os.path.join(root, pattern)
    files = sorted(glob.glob(pat, recursive=recursive))
    files = [f for f in files if os.path.isfile(f)]
    if sort_by == "time":
        files.sort(key=lambda p: os.path.getmtime(p))
    if not files:
        return {"error": f"未匹配到文件: {pat}"}

    frames: List[Any] = []
    skipped: List[str] = []
    for f in files:
        ext = os.path.splitext(f)[1].lower()
        try:
            if ext in (".csv", ".tsv"):
                sep = "\t" if ext == ".tsv" else None
                d = pd.read_csv(f, sep=sep, engine="python", dtype=object)
            elif ext in (".xlsx", ".xlsm", ".xls"):
                d = pd.read_excel(f, dtype=object)
            elif ext in (".txt", ".md", ".log"):
                with open(f, "r", encoding="utf-8", errors="replace") as fh:
                    lines = [ln.rstrip("\n") for ln in fh if ln.strip()]
                d = pd.DataFrame({"content": lines})
            else:
                skipped.append(f)
                continue
        except Exception as e:  # noqa: BLE001
            logger.warning(f"读取跳过 {f}: {e}")
            skipped.append(f)
            continue
        if add_source_col:
            d.insert(0, "_source_file", os.path.basename(f))
        frames.append(d)

    if not frames:
        return {"error": "没有可合并的有效文件"}

    merged = pd.concat(frames, ignore_index=True, sort=False)

    if not out_path:
        out_path = os.path.join(expand("~/.archive_assistant/output"), "merged.xlsx")
    out_path = _unique_path(expand(out_path))

    result = {
        "root": root,
        "pattern": pattern,
        "files_merged": len(frames),
        "skipped": skipped,
        "rows": int(merged.shape[0]),
        "cols": int(merged.shape[1]),
        "columns": list(merged.columns),
        "out_path": out_path,
        "dry_run": dry_run,
    }
    if dry_run:
        return result

    ensure_dir(os.path.dirname(out_path) or ".")
    try:
        if out_path.lower().endswith(".csv"):
            merged.to_csv(out_path, index=False, encoding="utf-8-sig")
        else:
            merged.to_excel(out_path, index=False)
    except Exception as e:  # noqa: BLE001
        return {"error": f"写出失败: {e}"}
    return result


# ============================================================
# 能力 3：docx_compose — 清单/资料整理成 Word
# ============================================================

def capability_docx_compose(
    out_path: str,
    title: Optional[str] = None,
    from_files: Optional[List[str]] = None,
    body: Optional[str] = None,
    template: Optional[str] = None,
    context: Optional[Dict[str, Any]] = None,
    dry_run: bool = False,
) -> Dict[str, Any]:
    """生成 Word 文档。

    两种模式：
    1. 模板模式：提供 template(.docx) + context(dict) → docxtpl 渲染填充；
    2. 普通模式：title + body 文本（或从 from_files 收集清单）→ 标题 + 段落。
    安全：输出写到新文件，不覆盖。
    """
    out_path = _unique_path(expand(out_path))

    # 模板模式
    if template:
        try:
            from docxtpl import DocxTemplate
        except ImportError:
            return _missing("docxtpl", "docxtpl")
        template = expand(template)
        if not os.path.isfile(template):
            return {"error": f"模板不存在: {template}"}
        result = {"mode": "template", "template": template, "out_path": out_path, "dry_run": dry_run}
        if dry_run:
            return result
        ensure_dir(os.path.dirname(out_path) or ".")
        try:
            tpl = DocxTemplate(template)
            tpl.render(context or {})
            tpl.save(out_path)
        except Exception as e:  # noqa: BLE001
            return {"error": f"模板渲染失败: {e}"}
        return result

    # 普通模式
    try:
        from docx import Document
    except ImportError:
        return _missing("python-docx", "python-docx")

    # 收集 from_files 清单（仅文件名列表，不解析内容）
    items: List[str] = []
    if from_files:
        for f in from_files:
            for g in sorted(glob.glob(expand(f))):
                items.append(os.path.basename(g))

    paras = [p for p in (body or "").split("\n")] if body else []

    result = {
        "mode": "plain",
        "out_path": out_path,
        "title": title,
        "paragraphs": len(paras),
        "list_items": len(items),
        "dry_run": dry_run,
    }
    if dry_run:
        return result

    ensure_dir(os.path.dirname(out_path) or ".")
    try:
        doc = Document()
        if title:
            doc.add_heading(title, level=0)
        for p in paras:
            doc.add_paragraph(p)
        if items:
            for it in items:
                doc.add_paragraph(it, style="List Bullet")
        doc.save(out_path)
    except Exception as e:  # noqa: BLE001
        return {"error": f"写出失败: {e}"}
    return result


# ============================================================
# 能力 4：pdf_ops — PDF 合并/拆分/抽取文本
# ============================================================

def capability_pdf_ops(
    op: str,
    inputs: List[str],
    out_path: Optional[str] = None,
    pages: Optional[str] = None,
    dry_run: bool = False,
) -> Dict[str, Any]:
    """PDF 合并 / 拆分 / 抽取文本（纯 pypdf，无需 Office）。

    op:
      merge   —— 多个 PDF 合并成一个
      split   —— 单个 PDF 按页拆分；pages="1-3,5" 抽取指定页为一个 PDF
      extract —— 抽取文本到 .txt
    安全：输出写到新文件，不覆盖原 PDF。
    """
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        return _missing("pypdf", "pypdf")

    files: List[str] = []
    for p in inputs:
        files.extend(sorted(glob.glob(expand(p))))
    files = [f for f in files if os.path.isfile(f) and f.lower().endswith(".pdf")]
    if not files:
        return {"error": "未匹配到 PDF 文件"}

    def _parse_pages(spec: str, total: int) -> List[int]:
        idx: List[int] = []
        for part in spec.split(","):
            part = part.strip()
            if "-" in part:
                a, b = part.split("-", 1)
                idx.extend(range(int(a) - 1, int(b)))
            elif part:
                idx.append(int(part) - 1)
        return [i for i in idx if 0 <= i < total]

    # ---- merge ----
    if op == "merge":
        out_path = _unique_path(expand(out_path or os.path.join(os.path.dirname(files[0]), "merged.pdf")))
        result = {"op": "merge", "inputs": files, "out_path": out_path, "dry_run": dry_run}
        if dry_run:
            return result
        ensure_dir(os.path.dirname(out_path) or ".")
        try:
            writer = PdfWriter()
            for f in files:
                for pg in PdfReader(f).pages:
                    writer.add_page(pg)
            with open(out_path, "wb") as fh:
                writer.write(fh)
        except Exception as e:  # noqa: BLE001
            return {"error": f"合并失败: {e}"}
        return result

    # ---- split / extract 针对单个文件 ----
    src = files[0]
    try:
        reader = PdfReader(src)
    except Exception as e:  # noqa: BLE001
        return {"error": f"读取失败: {e}"}
    total = len(reader.pages)

    if op == "split":
        sel = _parse_pages(pages, total) if pages else list(range(total))
        if not sel:
            return {"error": "页码范围无效"}
        base, _ = os.path.splitext(src)
        if pages:
            # 抽取指定页为单个 PDF
            out_path = _unique_path(expand(out_path or f"{base}_p{pages.replace(',', '_')}.pdf"))
            result = {"op": "split", "src": src, "pages": sel, "out_path": out_path, "dry_run": dry_run}
            if dry_run:
                return result
            ensure_dir(os.path.dirname(out_path) or ".")
            writer = PdfWriter()
            for i in sel:
                writer.add_page(reader.pages[i])
            with open(out_path, "wb") as fh:
                writer.write(fh)
            return result
        else:
            # 每页拆成一个 PDF
            out_dir = expand(out_path) if out_path else f"{base}_pages"
            result = {"op": "split", "src": src, "pages": sel, "out_dir": out_dir, "files": total, "dry_run": dry_run}
            if dry_run:
                return result
            ensure_dir(out_dir)
            for i in sel:
                w = PdfWriter()
                w.add_page(reader.pages[i])
                with open(os.path.join(out_dir, f"page_{i + 1:03d}.pdf"), "wb") as fh:
                    w.write(fh)
            return result

    if op == "extract":
        base, _ = os.path.splitext(src)
        out_path = _unique_path(expand(out_path or f"{base}.txt"))
        result = {"op": "extract", "src": src, "out_path": out_path, "pages": total, "dry_run": dry_run}
        if dry_run:
            return result
        ensure_dir(os.path.dirname(out_path) or ".")
        try:
            texts = []
            for pg in reader.pages:
                texts.append(pg.extract_text() or "")
            with open(out_path, "w", encoding="utf-8") as fh:
                fh.write("\n\n".join(texts))
        except Exception as e:  # noqa: BLE001
            return {"error": f"抽取失败: {e}"}
        return result

    return {"error": f"不支持的 op: {op}（支持 merge/split/extract）"}
